import logging
import os
import re
import subprocess
import tempfile
import zipfile
from typing import List, Optional, Dict, Any

import fitz  # PyMuPDF
import pandas as pd
from docx import Document as DocxDocument
from jnd_utils.log import init_logging

from src.models.unified_tender import ParsedContent, ParsedDocumentData, UnifiedTenderRecord
from src.utils.file_utils import FileType, get_file_type
from src.utils.http_client import HttpClient

logger = logging.getLogger(__name__)


def convert_doc_to_docx(path: str) -> str:
    """Convert a .doc file to .docx format using LibreOffice."""
    logger.info(f"Converting .doc file to .docx: {path}")
    if not path.lower().endswith(".doc"):
        logger.error(f"Invalid file extension for conversion: {path}")
        raise ValueError("Input file must be .doc")

    output_dir = tempfile.mkdtemp()
    try:
        logger.debug(f"Running LibreOffice conversion to output dir: {output_dir}")
        subprocess.run([
            "libreoffice", "--headless", "--convert-to", "docx", "--outdir", output_dir, path
        ], check=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)

        converted_name = os.path.splitext(os.path.basename(path))[0] + ".docx"
        converted_path = os.path.join(output_dir, converted_name)

        if not os.path.exists(converted_path):
            logger.error(f"Conversion failed - output file not found: {converted_path}")
            raise FileNotFoundError("Conversion failed, .docx not found.")

        logger.info(f"Successfully converted to: {converted_path}")
        return converted_path

    except subprocess.CalledProcessError as e:
        logger.error(f"LibreOffice conversion failed: {e}")
        raise RuntimeError(f"LibreOffice conversion failed: {e}")


class DocumentsParser:
    """
    Parser for various document types with metadata extraction.
    Integrated with UnifiedTenderRecord for seamless document processing.
    """

    def __init__(self,
                 document_infos: List[Dict[str, Any]],
                 http_client: Optional[HttpClient] = None,
                 tender_record: Optional['UnifiedTenderRecord'] = None):
        """
        Initialize the document parser.

        Args:
            document_infos: List of document information dictionaries.
                           Each dict should contain 'id', 'file', and 'download_link' keys.
            http_client: Optional HTTP client for downloading files
            tender_record: Optional UnifiedTenderRecord to directly update with parsed documents
        """
        logger.info("Initializing DocumentsParser")
        if not isinstance(document_infos, list):
            logger.error("Invalid document_infos type")
            raise TypeError("document_infos must be a list")

        self.documents_data: List[ParsedDocumentData] = []
        self.document_infos = document_infos
        self.http_client = http_client or HttpClient()
        self.tender_record = tender_record
        self._process_all_documents()

    def _is_url(self, path: str) -> bool:
        """Check if a path is a URL."""
        return path.startswith(('http://', 'https://'))

    def _download_file(self, name: str, url: str) -> str:
        """
        Download a file from a URL using the HTTP client.
        
        Args:
            name: The name of the file, used for file type detection
            url: The URL to download from
        
        Returns:
            Path to the downloaded file
        
        Raises:
            RuntimeError: If download fails
        """
        logger.info(f"Downloading file from URL: {url}")
        file_path = self.http_client.download_file(url, file_name=name)
        if not file_path:
            logger.error(f"Failed to download file: {url}")
            raise RuntimeError(f"Failed to download file: {url}")
        return file_path

    @staticmethod
    def _split_section_text(text: str) -> List[str]:
        """Split text into paragraphs."""
        paragraphs = re.split(r'\n{2,}|\r\n{2,}', text)
        return [p.strip() for p in paragraphs if len(p.strip()) > 30]

    @staticmethod
    def _handle_parsing_error(error_msg: str) -> ParsedContent:
        """Create a ParsedContent object for parsing errors."""
        logger.error(f"Parsing error: {error_msg}")
        return ParsedContent(
            preview=error_msg,
            full_text=""
        )

    def _parse_pdf(self, path: str) -> ParsedContent:
        """Parse a PDF file and extract text content."""
        logger.info(f"Parsing PDF file: {path}")
        if not os.path.exists(path):
            return self._handle_parsing_error("File not found")

        try:
            with fitz.open(path) as doc:
                text_blocks = []
                for page_num, page in enumerate(doc, 1):
                    logger.debug(f"Processing page {page_num}")
                    blocks = page.get_text("blocks")
                    for block in blocks:
                        if isinstance(block, tuple) and len(block) > 4:
                            text_blocks.append(block[4])

                full_text = "\n".join(text_blocks)
                preview = "\n".join(text_blocks[:5])
                logger.info(f"Successfully parsed PDF with {len(text_blocks)} blocks")

                return ParsedContent(
                    preview=preview,
                    full_text=full_text
                )

        except Exception as e:
            return self._handle_parsing_error(f"Error parsing PDF: {str(e)}")

    def _parse_docx(self, path: str) -> ParsedContent:
        """Parse a DOCX file and extract text content."""
        logger.info(f"Parsing DOCX file: {path}")
        try:
            doc = DocxDocument(path)
        except Exception as e:
            return self._handle_parsing_error(f"Failed to open DOCX file: {str(e)}")

        text_blocks = []
        try:
            logger.debug("Processing paragraphs")
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text = paragraph.text
                    try:
                        if paragraph.style and paragraph.style.name and \
                                paragraph.style.name.startswith('Heading'):
                            text = f"\n{text}\n"
                    except AttributeError:
                        pass  # Handle cases where style information is not available
                    text_blocks.append(text)

            logger.debug("Processing tables")
            for table in doc.tables:
                for row in table.rows:
                    # Preserve empty cells with space to maintain structure
                    row_text = ' | '.join(cell.text.strip() if cell.text.strip()
                                          else ' ' for cell in row.cells)
                    if any(cell.text.strip() for cell in row.cells):  # Add row if it has any content
                        text_blocks.append(row_text)

            full_text = "\n".join(text_blocks)
            preview = "\n".join(text_blocks[:5])
            logger.info(f"Successfully parsed DOCX with {len(text_blocks)} blocks")

            return ParsedContent(
                preview=preview,
                full_text=full_text
            )

        except Exception as e:
            return self._handle_parsing_error(f"Error parsing DOCX content: {str(e)}")

    def _parse_excel_csv(self, path: str) -> ParsedContent:
        """Parse an Excel or CSV file and extract text content."""
        logger.info(f"Parsing Excel/CSV file: {path}")
        try:
            df = pd.read_excel(path) if path.endswith((".xls", ".xlsx")) else pd.read_csv(path)

            if df.empty:
                logger.warning("Empty file detected")
                return self._handle_parsing_error("Empty file")

            full_text = df.to_string()
            preview = df.head().to_string()
            logger.info(f"Successfully parsed file with {len(df)} rows")

            return ParsedContent(
                preview=preview,
                full_text=full_text
            )
        except Exception as e:
            return self._handle_parsing_error(f"Failed to parse: {e}")

    def _parse_file(self, path: str, file_ext: str) -> ParsedContent:
        """Parse a file based on its type."""
        logger.info(f"Parsing file: {path} with extension {file_ext}")
        file_type = get_file_type(file_ext)

        if file_type == FileType.PDF:
            return self._parse_pdf(path)
        elif file_type == FileType.DOCX:
            return self._parse_docx(path)
        elif file_type == FileType.DOC:
            try:
                docx_path = convert_doc_to_docx(path)
                return self._parse_docx(docx_path)
            except Exception as e:
                return self._handle_parsing_error(f"Failed to convert .doc: {e}")
        elif file_type in [FileType.EXCEL, FileType.CSV]:
            return self._parse_excel_csv(path)
        else:
            return self._handle_parsing_error("Unsupported file type")

    def _unpack_zip(self, zip_path: str, doc_id: str) -> List[ParsedDocumentData]:
        """Extract and parse files from a ZIP archive."""
        logger.info(f"Unpacking ZIP file: {zip_path}")
        extracted = []

        # Create a more persistent directory for ZIP contents
        zip_extract_dir = os.path.join(os.path.dirname(zip_path), f"{doc_id}_extracted")
        os.makedirs(zip_extract_dir, exist_ok=True)

        try:
            with zipfile.ZipFile(zip_path, 'r') as zip_ref:
                logger.debug(f"Extracting ZIP to directory: {zip_extract_dir}")
                file_list = zip_ref.namelist()
                logger.debug(f"Found {len(file_list)} files in ZIP archive: {file_list}")

                # Filter out directories and hidden files
                valid_files = [f for f in file_list if not f.endswith('/') and not f.startswith('.')]
                zip_ref.extractall(zip_extract_dir)

                count = 1
                for file_name in valid_files:
                    full_path = os.path.join(zip_extract_dir, file_name)

                    # Skip if file doesn't exist (could be a directory entry)
                    if not os.path.isfile(full_path):
                        continue

                    ext = os.path.splitext(file_name)[-1].lower()
                    file_type = get_file_type(ext)

                    # Skip unsupported file types
                    if file_type == FileType.UNSUPPORTED:
                        logger.debug(f"Skipping unsupported file type: {file_name}")
                        continue

                    logger.debug(f"Processing ZIP content file {count}: {file_name}")

                    try:
                        parsed = self._parse_file(full_path, ext)

                        # Use more descriptive ID for ZIP contents
                        zip_content_id = f"{doc_id}_zip_{count}_{os.path.splitext(file_name)[0]}"

                        extracted.append(ParsedDocumentData(
                            id=zip_content_id,
                            name=file_name,
                            type=ext[1:] if ext.startswith('.') else ext,
                            path=full_path,
                            url=None,  # ZIP contents don't have URLs
                            preview=parsed.preview,
                            full_text=parsed.full_text
                        ))
                        count += 1

                    except Exception as e:
                        logger.error(f"Error parsing ZIP content file {file_name}: {e}")
                        # Create error metadata for failed ZIP content
                        error_doc = ParsedDocumentData(
                            id=f"{doc_id}_zip_{count}_error",
                            name=file_name,
                            type=ext[1:] if ext.startswith('.') else ext,
                            path=full_path,
                            url=None,
                            preview=f"Error parsing file: {str(e)}",
                            full_text=f"Failed to parse ZIP content file {file_name}: {str(e)}"
                        )
                        extracted.append(error_doc)
                        count += 1

        except Exception as e:
            logger.error(f"Error unpacking ZIP file {zip_path}: {e}")
            # Return error document for the entire ZIP
            error_doc = ParsedDocumentData(
                id=f"{doc_id}_zip_error",
                name=os.path.basename(zip_path),
                type="zip",
                path=zip_path,
                url=None,
                preview=f"Error unpacking ZIP: {str(e)}",
                full_text=f"Failed to unpack ZIP file: {str(e)}"
            )
            return [error_doc]

        logger.info(f"Successfully extracted and parsed {len(extracted)} files from ZIP")
        return extracted

    @staticmethod
    def _create_error_metadata(doc: Dict[str, Any], error: Exception) -> ParsedDocumentData:
        """Create metadata for a document that failed to process."""
        logger.error(f"Creating error metadata for document {doc.get('id', 'unknown')}: {error}")
        return ParsedDocumentData(
            id=doc.get('id', "unknown"),
            name=doc.get("file", "unknown"),
            url=doc.get("download_link", ""),
            type="error",
            path=None,
            preview=f"Error processing file: {error}",
            full_text=""
        )

    def _process_document(self, doc: Dict[str, Any]) -> None:
        """Process a single document."""
        doc_id = doc.get('id', 'unknown')
        logger.info(f"Processing document: {doc_id}")
        try:
            name = doc.get("file")
            url = doc.get("download_link")

            if not all([doc_id, name, url]):
                logger.error(f"Missing required document information for {doc_id}")
                raise ValueError("Missing required document information")

            path = self._download_file(name, url)
            ext = os.path.splitext(name)[-1].lower()
            file_type = get_file_type(ext)

            if file_type == FileType.ZIP:
                # Handle ZIP files specially
                logger.info(f"Processing ZIP file: {name}")
                zip_contents = self._unpack_zip(path, doc_id)

                # Add the ZIP file itself as a document
                zip_doc = ParsedDocumentData(
                    id=doc_id,
                    name=name,
                    url=url,
                    type="zip",
                    path=path,
                    preview=f"ZIP archive containing {len(zip_contents)} files",
                    full_text=f"ZIP archive: {name} - Contains {len(zip_contents)} extractable files"
                )
                self.documents_data.append(zip_doc)

                # Add ZIP file directly to tender record if available
                if self.tender_record:
                    self.tender_record.add_parsed_document(zip_doc)

                # Add all extracted contents
                self.documents_data.extend(zip_contents)

                # Add extracted contents to tender record if available
                if self.tender_record:
                    for content_doc in zip_contents:
                        self.tender_record.add_parsed_document(content_doc)

            elif file_type == FileType.UNSUPPORTED:
                logger.warning(f"Unsupported file type for {name}: {ext}")
                # Still create a document entry for tracking
                unsupported_doc = ParsedDocumentData(
                    id=doc_id,
                    name=name,
                    url=url,
                    type=ext[1:] if ext.startswith('.') else ext,
                    path=path,
                    preview=f"Unsupported file type: {ext}",
                    full_text=f"File {name} has unsupported format: {ext}"
                )
                self.documents_data.append(unsupported_doc)

                # Add to tender record if available
                if self.tender_record:
                    self.tender_record.add_parsed_document(unsupported_doc)
            else:
                # Handle regular files
                parsed = self._parse_file(path, ext)
                regular_doc = ParsedDocumentData(
                    id=doc_id,
                    name=name,
                    url=url,
                    type=ext[1:] if ext.startswith('.') else ext,
                    path=path,
                    preview=parsed.preview,
                    full_text=parsed.full_text
                )
                self.documents_data.append(regular_doc)

                # Add to tender record if available
                if self.tender_record:
                    self.tender_record.add_parsed_document(regular_doc)

            logger.info(f"Successfully processed document: {doc_id}")

        except Exception as e:
            logger.error(f"Error processing document {doc_id}: {e}")
            error_doc = self._create_error_metadata(doc, e)
            self.documents_data.append(error_doc)

            # Add error document to tender record if available
            if self.tender_record:
                self.tender_record.add_parsed_document(error_doc)
                self.tender_record.add_processing_error(f"Document parsing error for {doc_id}: {str(e)}")

    def _process_all_documents(self) -> None:
        """Process all documents in the document_infos list."""
        logger.info(f"Processing {len(self.document_infos)} documents")
        for doc in self.document_infos:
            self._process_document(doc)
        logger.info("Completed processing all documents")

        # Update tender record processing stage if available
        if self.tender_record:
            from src.models.unified_tender import ProcessingStage
            self.tender_record.processing_stage = ProcessingStage.DOCUMENTS_PARSED

    def get_documents_data(self) -> List[ParsedDocumentData]:
        """Get metadata for all processed documents."""
        return self.documents_data

    def get_full_data(self, name: str) -> Optional[ParsedDocumentData]:
        """Get full document data by name."""
        logger.debug(f"Searching for document: {name}")
        for doc in self.documents_data:
            if doc.name == name:
                logger.debug(f"Found document: {name}")
                return doc
        logger.debug(f"Document not found: {name}")
        return None

    @classmethod
    def parse_documents_for_tender(cls,
                                   tender_record: UnifiedTenderRecord,
                                   http_client: Optional[HttpClient] = None) -> 'DocumentsParser':
        """
        Class method to parse documents directly for a UnifiedTenderRecord.
        
        Args:
            tender_record: The tender record to parse documents for
            http_client: Optional HTTP client for downloading files
            
        Returns:
            DocumentsParser instance with processed documents
        """
        logger.info(f"Parsing documents for tender: {tender_record.tender_id}")

        if not tender_record.document_infos:
            logger.warning(f"No document infos available for tender: {tender_record.tender_id}")
            return cls([], http_client=http_client, tender_record=tender_record)

        return cls(
            document_infos=tender_record.document_infos,
            http_client=http_client,
            tender_record=tender_record
        )


if __name__ == '__main__':
    init_logging()
    logger.info("Starting document parser")
    docs_items = [
        {
            'download_link': 'https://nen.nipez.cz/file?id=2845592172',
            'file': 'test_pdf_file.pdf',
            'id': '2845592172',
        }
    ]
    parser = DocumentsParser(docs_items)
    print(parser.get_documents_data())
